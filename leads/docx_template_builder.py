"""Gestaltete Word-Startervorlagen im Corporate Design (AutoPark Grün)."""

from __future__ import annotations

import io
from pathlib import Path

from django.conf import settings
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import DealerDocumentTemplate

BRAND_GREEN = RGBColor(0x66, 0xBB, 0x46)
BRAND_GREEN_DARK = RGBColor(0x57, 0xA6, 0x3A)
BRAND_BLACK = RGBColor(0x00, 0x00, 0x00)
TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)
TEXT_BODY = RGBColor(0x22, 0x22, 0x22)
FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def _add_styled_run(paragraph, text: str, *, bold: bool = False, size: int = 11, color=None, italic: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _add_placeholder_paragraph(doc, text: str, *, bold: bool = False, size: int = 11, align=None, after: int = 6):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    _add_styled_run(paragraph, text, bold=bold, size=size, color=TEXT_BODY)
    _set_paragraph_spacing(paragraph, after=after)
    return paragraph


def _add_section_heading(doc, title: str) -> None:
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, before=10, after=4)
    _add_styled_run(paragraph, title, bold=True, size=12, color=BRAND_GREEN_DARK)


def _add_body(doc, text: str, *, after: int = 6) -> None:
    _add_placeholder_paragraph(doc, text, after=after)


def _add_legal_note(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, before=8, after=8)
    _add_styled_run(paragraph, text, size=9, color=TEXT_MUTED, italic=True)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def _find_logo_path() -> Path | None:
    logo_candidates = [
        Path(settings.BASE_DIR) / "static" / settings.SITE_LOGO_PATH.lstrip("/"),
        Path(settings.BASE_DIR) / "static" / "img" / "logo.png",
        Path(settings.BASE_DIR) / "static" / "img" / "logo.jpg",
    ]
    for logo_path in logo_candidates:
        if logo_path.is_file():
            return logo_path
    return None


def _add_brand_to_cell(cell) -> None:
    logo_path = _find_logo_path()
    logo_para = cell.paragraphs[0]
    if logo_path:
        run = logo_para.add_run()
        run.add_picture(str(logo_path), width=Cm(4.5))
        return
    _add_styled_run(logo_para, settings.SITE_BRAND_NAME, bold=True, size=18, color=BRAND_GREEN)
    _add_styled_run(logo_para, "\n", size=6)
    _add_styled_run(logo_para, settings.SITE_COMPANY_NAME, size=9, color=TEXT_MUTED)


def _add_corporate_header(doc: Document, document_title: str, subtitle: str = "") -> None:
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]

    _add_brand_to_cell(left_cell)

    meta_para = right_cell.paragraphs[0]
    meta_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(meta_para, document_title + "\n", bold=True, size=16, color=BRAND_BLACK)
    if subtitle:
        _add_styled_run(meta_para, subtitle + "\n", size=10, color=TEXT_MUTED)
    _add_styled_run(meta_para, "{{ firma_name }}\n", size=9, color=TEXT_BODY)
    _add_styled_run(meta_para, "{{ firma_adresse }}\n", size=9, color=TEXT_BODY)
    _add_styled_run(meta_para, "Tel. {{ firma_telefon }} · {{ firma_email }}", size=9, color=TEXT_BODY)

    divider = doc.add_table(rows=1, cols=1)
    divider_cell = divider.rows[0].cells[0]
    _set_cell_shading(divider_cell, "66BB46")
    divider_cell.paragraphs[0].add_run(" ")
    divider_cell.height = Cm(0.12)

    doc.add_paragraph()


def _add_corporate_footer(doc: Document) -> None:
    doc.add_paragraph()
    footer_table = doc.add_table(rows=1, cols=1)
    footer_cell = footer_table.rows[0].cells[0]
    _set_cell_shading(footer_cell, "E8F4E1")
    footer_para = footer_cell.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(footer_para, "{{ firma_name }} · {{ firma_adresse }}\n", size=8, color=TEXT_MUTED)
    _add_styled_run(
        footer_para,
        "Tel. {{ firma_telefon }} · {{ firma_email }} · Lead-Nr. {{ lead_id }}",
        size=8,
        color=TEXT_MUTED,
    )


def _add_signature_block(doc, left_label: str, right_label: str) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, label in enumerate((left_label, right_label)):
        cell = table.rows[0].cells[col]
        para = cell.paragraphs[0]
        _add_styled_run(para, "Ort, Datum", size=9, color=TEXT_MUTED)

        line_cell = table.rows[1].cells[col]
        line_para = line_cell.paragraphs[0]
        _add_styled_run(line_para, "_______________________________", size=10, color=TEXT_BODY)

        name_cell = table.rows[2].cells[col]
        name_para = name_cell.paragraphs[0]
        _add_styled_run(name_para, label, size=9, color=TEXT_MUTED)


def _build_purchase_contract(doc: Document) -> None:
    _configure_document(doc)
    _add_corporate_header(
        doc,
        "Kaufvertrag",
        "Fahrzeugankauf · Gebrauchtwagen",
    )

    _add_body(
        doc,
        "Zwischen dem nachstehend genannten Verkäufer und dem nachstehend genannten Käufer "
        "wird folgender Kaufvertrag über ein gebrauchtes Kraftfahrzeug geschlossen.",
    )

    parties = doc.add_table(rows=2, cols=2)
    parties.style = "Table Grid"
    parties.alignment = WD_TABLE_ALIGNMENT.CENTER

    seller_head = parties.rows[0].cells[0]
    _set_cell_shading(seller_head, "E8F4E1")
    _add_styled_run(seller_head.paragraphs[0], "Verkäufer (Privatperson)", bold=True, size=10, color=BRAND_GREEN_DARK)
    buyer_head = parties.rows[0].cells[1]
    _set_cell_shading(buyer_head, "E8F4E1")
    _add_styled_run(buyer_head.paragraphs[0], "Käufer (Händler)", bold=True, size=10, color=BRAND_GREEN_DARK)

    seller_body = parties.rows[1].cells[0]
    _add_styled_run(seller_body.paragraphs[0], "{{ verkaeufer_name }}\n", bold=True, size=11)
    _add_styled_run(seller_body.paragraphs[0], "PLZ {{ verkaeufer_plz }}\n", size=10, color=TEXT_BODY)
    _add_styled_run(seller_body.paragraphs[0], "Tel. {{ kunde_telefon }}\n", size=10, color=TEXT_BODY)
    _add_styled_run(seller_body.paragraphs[0], "E-Mail: {{ kunde_email }}", size=10, color=TEXT_BODY)

    buyer_body = parties.rows[1].cells[1]
    _add_styled_run(buyer_body.paragraphs[0], "{{ kaeufer_name }}\n", bold=True, size=11)
    _add_styled_run(buyer_body.paragraphs[0], "{{ kaeufer_adresse }}\n", size=10, color=TEXT_BODY)
    _add_styled_run(buyer_body.paragraphs[0], "Tel. {{ firma_telefon }}\n", size=10, color=TEXT_BODY)
    _add_styled_run(buyer_body.paragraphs[0], "E-Mail: {{ firma_email }}", size=10, color=TEXT_BODY)

    doc.add_paragraph()

    _add_section_heading(doc, "§ 1 Vertragsgegenstand")
    _add_body(doc, "{{ vertragsgegenstand }}")

    vehicle_box = doc.add_table(rows=1, cols=1)
    vehicle_cell = vehicle_box.rows[0].cells[0]
    _set_cell_shading(vehicle_cell, "F7FBF5")
    vehicle_para = vehicle_cell.paragraphs[0]
    _add_styled_run(vehicle_para, "Fahrzeug:\n", bold=True, size=10, color=BRAND_GREEN_DARK)
    _add_styled_run(vehicle_para, "{{ auto }}\n\n", bold=True, size=11)
    _add_styled_run(vehicle_para, "{{ auto_beschreibung_lang }}", size=10, color=TEXT_BODY)

    _add_section_heading(doc, "§ 2 Kaufpreis und Zahlungsweise")
    _add_body(
        doc,
        "Der vereinbarte Kaufpreis für das Fahrzeug beträgt {{ preisvorstellung }} EUR "
        "(in Worten: _____________________________________________ Euro).",
    )
    _add_body(
        doc,
        "Der Kaufpreis wird nach Übergabe des Fahrzeugs und der vollständigen Unterlagen "
        "(Fahrzeugbrief, Fahrzeugschein, Schlüssel, TÜV-/AU-Nachweis soweit vorhanden) "
        "durch Überweisung auf das vom Verkäufer benannte Konto oder in bar gegen Quittung geleistet.",
    )

    _add_section_heading(doc, "§ 3 Übergabe und Eigentumsübergang")
    _add_body(
        doc,
        "Die Übergabe des Fahrzeugs erfolgt am _________________________ in "
        "_________________________________________ (Ort). Mit Übergabe und vollständiger "
        "Zahlung des Kaufpreises geht das Eigentum am Fahrzeug auf den Käufer über.",
    )
    _add_body(
        doc,
        "Der Verkäufer übergibt das Fahrzeug im Zustand: {{ zustand }}. "
        "Kilometerstand bei Übergabe: {{ kilometerstand }} km. "
        "Angemeldet: {{ angemeldet }}. TÜV gültig bis: {{ tuv_bis }}.",
    )

    _add_section_heading(doc, "§ 4 Gewährleistung")
    _add_body(
        doc,
        "Der Verkäufer verkauft als Privatperson. Es wird ausdrücklich vereinbart, dass "
        "die gesetzliche Gewährleistung für Sachmängel – soweit gesetzlich zulässig – "
        "ausgeschlossen ist, sofern nicht arglistig verschwiegene Mängel vorliegen.",
    )
    _add_body(
        doc,
        "Der Käufer hat das Fahrzeug besichtigt und Probe gefahren. Bekannte Mängel, "
        "Unfallschäden und Vorschäden sind in der nachfolgenden Erklärung aufgeführt "
        "bzw. als „nicht bekannt“ zu bezeichnen.",
    )

    _add_section_heading(doc, "§ 5 Erklärungen des Verkäufers")
    _add_body(doc, "Der Verkäufer erklärt gegenüber dem Käufer verbindlich:")
    for item in (
        "Er ist alleiniger und uneingeschränkter Eigentümer des Fahrzeugs.",
        "Das Fahrzeug ist frei von Rechten und Ansprüchen Dritter (Pfandrecht, Sicherungseigentum o. Ä.).",
        "Es bestehen keine offenen Kredit- oder Leasingverbindlichkeiten am Fahrzeug.",
        "Der angegebene Kilometerstand entspricht nach bestem Wissen und Gewissen dem tatsächlichen Stand.",
        "Unfallfreiheit / Vorschäden: ________________________________________________",
        "Anzahl Schlüssel bei Übergabe: ______",
    ):
        _add_body(doc, f"• {item}", after=3)

    _add_section_heading(doc, "§ 6 Datenschutz und Einwilligung")
    _add_body(
        doc,
        "Der Verkäufer willigt ein, dass seine angegebenen Kontaktdaten (Name, Adresse/PLZ, "
        "Telefon, E-Mail) zum Zwecke der Vertragsabwicklung, Fahrzeugübernahme und "
        "Kommunikation durch den Käufer verarbeitet werden. Eine Weitergabe an unbefugte "
        "Dritte erfolgt nicht.",
    )

    _add_section_heading(doc, "§ 7 Schlussbestimmungen")
    _add_body(
        doc,
        "Änderungen und Ergänzungen dieses Vertrages bedürfen der Schriftform. "
        "Sollten einzelne Bestimmungen unwirksam sein, bleibt die Wirksamkeit der "
        "übrigen Regelungen unberührt (Salvatorische Klausel).",
    )
    _add_body(doc, "Es gilt das Recht der Bundesrepublik Deutschland. Gerichtsstand ist – soweit zulässig – Essen.")

    _add_signature_block(
        doc,
        "{{ verkaeufer_name }}\n(Verkäufer)",
        "{{ kaeufer_name }}\n(Käufer)",
    )

    _add_legal_note(
        doc,
        "Hinweis: Diese Vorlage dient als Muster für den Fahrzeugankauf von Privatpersonen. "
        "Für rechtsverbindliche Einzelfälle wird empfohlen, die Vorlage anwaltlich prüfen zu lassen.",
    )
    _add_corporate_footer(doc)


def _build_invoice_email(doc: Document) -> None:
    _configure_document(doc)
    _add_corporate_header(
        doc,
        "Ankaufsangebot",
        "Rechnung / E-Mail-Vorlage",
    )

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = meta_table.rows[0].cells[0]
    _add_styled_run(left.paragraphs[0], "An:", size=9, color=TEXT_MUTED)
    _add_styled_run(left.paragraphs[0], "\n{{ kunde_name }}\n", bold=True, size=11)
    _add_styled_run(left.paragraphs[0], "PLZ {{ kunde_plz }}\n", size=10)
    _add_styled_run(left.paragraphs[0], "{{ kunde_email }}\n", size=10)
    _add_styled_run(left.paragraphs[0], "Tel. {{ kunde_telefon }}", size=10)

    right = meta_table.rows[0].cells[1]
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_styled_run(right.paragraphs[0], "Datum: {{ datum_heute }}\n", size=10)
    _add_styled_run(right.paragraphs[0], "Angebots-Nr.: APG-{{ lead_id }}\n", size=10, bold=True, color=BRAND_GREEN_DARK)
    _add_styled_run(right.paragraphs[0], "Lead-Nr.: {{ lead_id }}", size=10, color=TEXT_MUTED)

    doc.add_paragraph()
    subject = doc.add_paragraph()
    _set_paragraph_spacing(subject, before=4, after=8)
    _add_styled_run(subject, "Betreff: ", bold=True, size=11)
    _add_styled_run(subject, "Ankaufsangebot für Ihr Fahrzeug {{ auto }}", size=11, color=BRAND_GREEN_DARK)

    _add_body(doc, "Sehr geehrte/r {{ kunde_name }},")
    _add_body(
        doc,
        "vielen Dank für Ihre Anfrage über unser Online-Formular. Bezugnehmend auf Ihr Fahrzeug "
        "{{ auto_beschreibung }} unterbreiten wir Ihnen folgendes verbindliches Ankaufsangebot "
        "unter den nachstehenden Bedingungen.",
    )

    offer_table = doc.add_table(rows=5, cols=2)
    offer_table.style = "Table Grid"
    offer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ("Position", "Details")
    for idx, label in enumerate(headers):
        cell = offer_table.rows[0].cells[idx]
        _set_cell_shading(cell, "66BB46")
        para = cell.paragraphs[0]
        _add_styled_run(para, label, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    rows = (
        ("Fahrzeug", "{{ auto }}"),
        ("Technische Daten", "{{ auto_beschreibung_lang }}"),
        ("Merkmale / Extras", "Merkmale: {{ merkmale }}\nExtras: {{ extras }}"),
        ("Ankaufspreis", "{{ preisvorstellung }} EUR"),
    )
    for row_idx, (label, value) in enumerate(rows, start=1):
        label_cell = offer_table.rows[row_idx].cells[0]
        _set_cell_shading(label_cell, "E8F4E1")
        _add_styled_run(label_cell.paragraphs[0], label, bold=True, size=10, color=BRAND_GREEN_DARK)
        value_cell = offer_table.rows[row_idx].cells[1]
        _add_styled_run(value_cell.paragraphs[0], value, size=10, color=TEXT_BODY)

    doc.add_paragraph()
    _add_section_heading(doc, "E-Mail-Text (zum Kopieren)")
    email_box = doc.add_table(rows=1, cols=1)
    email_cell = email_box.rows[0].cells[0]
    _set_cell_shading(email_cell, "F7FBF5")
    email_para = email_cell.paragraphs[0]
    email_text = (
        "Betreff: Ihr Ankaufsangebot für {{ auto }}\n\n"
        "Sehr geehrte/r {{ kunde_name }},\n\n"
        "wir können Ihr Fahrzeug {{ auto }} zum Preis von {{ preisvorstellung }} EUR ankaufen.\n\n"
        "Fahrzeugdetails:\n{{ auto_beschreibung }}\n\n"
        "Bei Interesse melden Sie sich bitte telefonisch unter {{ firma_telefon }} "
        "oder per E-Mail an {{ firma_email }}.\n\n"
        "Mit freundlichen Grüßen\n"
        "{{ firma_name }}\n"
        "{{ firma_adresse }}"
    )
    _add_styled_run(email_para, email_text, size=10, color=TEXT_BODY)

    _add_section_heading(doc, "Rechnung / Zahlungsbeleg (Entwurf)")
    _add_body(
        doc,
        "Die nachfolgende Rechnung dokumentiert den vereinbarten Ankaufspreis. "
        "Bei Zahlung durch den Käufer an den Verkäufer dient sie als Begleitdokument.",
        after=8,
    )

    invoice_table = doc.add_table(rows=6, cols=2)
    invoice_table.style = "Table Grid"
    invoice_rows = (
        ("Rechnungs-Nr.", "APG-{{ lead_id }}-{{ datum_heute }}"),
        ("Leistungsdatum", "{{ datum_heute }}"),
        ("Verkäufer", "{{ verkaeufer_name }}, PLZ {{ verkaeufer_plz }}"),
        ("Käufer / Rechnungssteller", "{{ kaeufer_name }}, {{ kaeufer_adresse }}"),
        ("Leistungsbeschreibung", "Ankauf gebrauchtes Kraftfahrzeug: {{ auto }}"),
        ("Bruttobetrag", "{{ preisvorstellung }} EUR"),
    )
    for idx, (label, value) in enumerate(invoice_rows):
        label_cell = invoice_table.rows[idx].cells[0]
        if idx == 0:
            _set_cell_shading(label_cell, "E8F4E1")
        _add_styled_run(label_cell.paragraphs[0], label, bold=True, size=10, color=TEXT_BODY)
        _add_styled_run(invoice_table.rows[idx].cells[1].paragraphs[0], value, size=10, color=TEXT_BODY)

    _add_legal_note(
        doc,
        "Hinweis: Bei Ankäufen von Privatpersonen kann die steuerliche Behandlung "
        "(z. B. Differenzbesteuerung nach §25a UStG) abweichen. Bitte Steuerberater konsultieren. "
        "Diese Vorlage ersetzt keine Rechts- oder Steuerberatung.",
    )
    _add_corporate_footer(doc)


def build_starter_template_docx(template_type: str) -> tuple[bytes, str]:
    doc = Document()

    if template_type == DealerDocumentTemplate.TYPE_PURCHASE_CONTRACT:
        _build_purchase_contract(doc)
        filename = "AutoParkGruen_Kaufvertrag_Vorlage.docx"
    elif template_type == DealerDocumentTemplate.TYPE_INVOICE_EMAIL:
        _build_invoice_email(doc)
        filename = "AutoParkGruen_Rechnung_Email_Vorlage.docx"
    else:
        raise ValueError(f"Unbekannter Vorlagentyp: {template_type}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), filename
